import os
import sys
import subprocess
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor

from docx2pdf import convert
from ebooklib import epub
from docx import Document

class DocxConverter:
    """
    Convert DOCX files to PDF, EPUB, and FB2 formats
    """
    
    def __init__(self, input_path, output_dir=None):
        """
        Initialize converter with input path and output directory
        """
        self.input_path = Path(input_path)
        
        # Validate input file exists and is a docx file
        if not self.input_path.exists():
            raise FileNotFoundError(f"Input file not found: {input_path}")
        
        if self.input_path.suffix.lower() != '.docx':
            raise ValueError(f"Input file must be a .docx file: {input_path}")
        
        # Set output directory
        if output_dir:
            self.output_dir = Path(output_dir)
            # Create output directory if it doesn't exist
            self.output_dir.mkdir(parents=True, exist_ok=True)
        else:
            self.output_dir = self.input_path.parent
        
        # Base filename without extension
        self.base_name = self.input_path.stem
        
        # Output paths
        self.pdf_path = self.output_dir / f"{self.base_name}.pdf"
        self.epub_path = self.output_dir / f"{self.base_name}.epub"
        self.fb2_path = self.output_dir / f"{self.base_name}.fb2"
    
    def convert_to_pdf(self):
        """
        Convert DOCX to PDF using docx2pdf
        """
        try:
            convert(str(self.input_path), str(self.pdf_path))
            print(f"PDF conversion complete: {self.pdf_path}")
            return str(self.pdf_path)
        except Exception as e:
            print(f"Error converting to PDF: {e}")
            raise
    
    def convert_to_epub(self):
        """
        Convert DOCX to EPUB using ebooklib
        """
        try:
            # Create a new EPUB book
            book = epub.EpubBook()
            
            # Set metadata
            book.set_identifier(f"id-{self.base_name}")
            book.set_title(self.base_name)
            book.set_language('en')
            book.add_author('Document Author')  # Could extract from document metadata
            
            # Extract text from DOCX
            doc = Document(self.input_path)
            full_text = "\n".join([para.text for para in doc.paragraphs])
            
            # Create chapter
            chapter = epub.EpubHtml(title=self.base_name, file_name='chapter.xhtml')
            chapter.content = f"<h1>{self.base_name}</h1>"
            
            # Add paragraphs
            for para_text in full_text.split("\n"):
                if para_text.strip():
                    chapter.content += f"<p>{para_text}</p>"
            
            # Add chapter to book
            book.add_item(chapter)
            
            # Define Table of Contents
            book.toc = [(epub.Section('Main'), [chapter])]
            
            # Add default NCX and Nav files
            book.add_item(epub.EpubNcx())
            book.add_item(epub.EpubNav())
            
            # Define CSS style
            style = """
            @namespace epub "http://www.idpf.org/2007/ops";
            body {
                font-family: Cambria, Liberation Serif, Bitstream Vera Serif, Georgia, Times, Times New Roman, serif;
            }
            h1 {
                text-align: center;
                margin-bottom: 1em;
            }
            """
            
            # Add CSS file
            nav_css = epub.EpubItem(
                uid="style_nav",
                file_name="style/nav.css",
                media_type="text/css",
                content=style
            )
            book.add_item(nav_css)
            
            # Create spine
            book.spine = ['nav', chapter]
            
            # Write EPUB file
            epub.write_epub(str(self.epub_path), book)
            
            print(f"EPUB conversion complete: {self.epub_path}")
            return str(self.epub_path)
        except Exception as e:
            print(f"Error converting to EPUB: {e}")
            raise
    
    def convert_to_fb2(self):
        """
        Convert DOCX to FB2 using fb2converter
        """
        try:
            # FB2 conversion is more complex, we'll use a temporary HTML file as an intermediate step
            html_path = self.output_dir / f"{self.base_name}.html"
            
            # First convert DOCX to HTML
            doc = Document(self.input_path)
            html_content = f"<!DOCTYPE html>\n<html>\n<head>\n<title>{self.base_name}</title>\n</head>\n<body>\n"
            html_content += f"<h1>{self.base_name}</h1>\n"
            
            for para in doc.paragraphs:
                if para.text.strip():
                    html_content += f"<p>{para.text}</p>\n"
            
            html_content += "</body>\n</html>"
            
            # Write HTML to temp file
            with open(html_path, 'w', encoding='utf-8') as html_file:
                html_file.write(html_content)
            
            # Use fb2converter to convert HTML to FB2
            # This assumes fb2converter is installed and available in the system path
            cmd = ['fb2converter', str(html_path), str(self.fb2_path)]
            
            try:
                subprocess.run(cmd, check=True)
                print(f"FB2 conversion complete: {self.fb2_path}")
            except FileNotFoundError:
                # Alternative method if fb2converter is not available
                # This is a simplified FB2 structure
                fb2_content = f"""<?xml version="1.0" encoding="UTF-8"?>
<FictionBook xmlns="http://www.gribuser.ru/xml/fictionbook/2.0">
    <description>
        <title-info>
            <genre>unspecified</genre>
            <author>
                <first-name>Document</first-name>
                <last-name>Author</last-name>
            </author>
            <book-title>{self.base_name}</book-title>
            <lang>en</lang>
        </title-info>
    </description>
    <body>
        <title><p>{self.base_name}</p></title>
"""
                # Add paragraphs
                for para in doc.paragraphs:
                    if para.text.strip():
                        fb2_content += f"        <p>{para.text}</p>\n"
                
                fb2_content += """    </body>
</FictionBook>"""
                
                # Write FB2 file
                with open(self.fb2_path, 'w', encoding='utf-8') as fb2_file:
                    fb2_file.write(fb2_content)
                
                print(f"FB2 conversion complete (alternative method): {self.fb2_path}")
            
            # Clean up temp HTML file
            if html_path.exists():
                os.remove(html_path)
                
            return str(self.fb2_path)
            
        except Exception as e:
            print(f"Error converting to FB2: {e}")
            raise
            
    def convert_all(self):
        """
        Convert DOCX to all supported formats
        """
        results = {}
        
        # Using ThreadPoolExecutor to run conversions in parallel
        with ThreadPoolExecutor(max_workers=3) as executor:
            # Submit all conversion tasks
            pdf_future = executor.submit(self.convert_to_pdf)
            epub_future = executor.submit(self.convert_to_epub)
            fb2_future = executor.submit(self.convert_to_fb2)
            
            # Get results
            try:
                results['pdf'] = pdf_future.result()
            except Exception as e:
                print(f"PDF conversion failed: {e}")
                
            try:
                results['epub'] = epub_future.result()
            except Exception as e:
                print(f"EPUB conversion failed: {e}")
                
            try:
                results['fb2'] = fb2_future.result()
            except Exception as e:
                print(f"FB2 conversion failed: {e}")
        
        print("All conversions completed")
        return results


def main():
    """
    Main function to handle command line arguments
    """
    if len(sys.argv) < 2:
        print("Usage: python docx_converter.py input.docx [output_directory]")
        # sys.exit(1)
    
    # input_path = sys.argv[1]
    input_path = "Frontier Shangri La.docx"
    output_dir = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        converter = DocxConverter(input_path, output_dir)
        results = converter.convert_all()
        print("\nConversion results:")
        for format_name, path in results.items():
            print(f"{format_name.upper()}: {path}")
    except Exception as e:
        print(f"Error: {e}")
        sys.exit(1)


main()