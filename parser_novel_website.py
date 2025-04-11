from selectolax.parser import HTMLParser
from docx import Document
import httpx
import asyncio
import random
import time


url = "https://ncode.syosetu.com/n6169dz/"
USER_AGENTS = [
    'Mozila/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4482.126 Safari/537.36',
    'Mozila/5.0 (Macintosh; Intel Mac OS X 10_11_5) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/960.0.4761.104 Safari/537.36',
    'Mozila/5.0 (X11; Linux x86_x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/50.0.4671.102 Safari/537.36',
    'Mozila/5.0 (iPhone; CPU iPhone OS 15_2 like Mac OS X) AppleWebKit/605.1.15 (KHTML, like Gecko) Version/15.2 Mobile/15E148 Safari/604.1',
]

headers = {
    "User-Agent": random.choice(USER_AGENTS),
    "Accept-Language": "en-Us,en;q=0.9",
    "Referer": "https://www.google.com/",
    "DNT": "1"
}


async def parse(count, part, parts_count, chapters, path):
    start = count * part + 1
    end = count * (part + 1) if part < parts_count - 1 else chapters + 1

    async with httpx.AsyncClient() as client:
        for i in range(start, end):
            res = await client.get(url + str(i), headers=headers)
            if res.status_code != 200:
                print(f"Failed with fetch: {url}{i}")
                continue

            tree = HTMLParser(res.text)

            title = tree.css_first("h1")
            paragraphs = tree.css("p")

            document = Document("source.docx")
            document.add_paragraph(title.text())

            for p in paragraphs:
                if "id" in p.attributes and "L" in p.attributes["id"]:
                    document.add_paragraph(p.text())
                    
            # document.save(f"Frontier Shangri La\\{i}.docx")
            document.save(f"{path}\\{i}.docx")
            print(f"Parsed {i}")
            document._body.clear_content()
            await asyncio.sleep(4)


# Path must be just folder name or path to folder - //String//, Example:  "Frontier Shangri La" or "folder_1\\folder_2\\Frontier Shangri La" 
async def multiple_parsing(chapters, number_of_threads, path):
    count = chapters // number_of_threads 
    tasks = [asyncio.create_task(parse(count, i, number_of_threads, chapters, path)) for i in range(number_of_threads)]
    await asyncio.gather(*tasks)

async def single_parsing(chapter, path, number_of_threads = 1):
    tasks = [asyncio.create_task(parse(chapter - 1, 1, number_of_threads, chapter + 1, path))]
    await asyncio.gather(*tasks)


# If need, implement nubmer of chapters/pages - find "p" with class="p-novel__number js-siori" and split to take last part 


# s = time.perf_counter()
# asyncio.run(multiple_parsing(939, 6))
# asyncio.run(single_parsing(939))
# e = time.perf_counter() - s
# print(f"Time: {e:0.6f} sec")

