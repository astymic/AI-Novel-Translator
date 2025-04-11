from google import genai
from google.genai import types
from docx import Document
from pathlib import Path
import json
import time
import re

with open('token.json', 'r') as file:
    data = json.load(file)
token = data["token"]

client = genai.Client(api_key=token)

# Balanced version
# prompt = """Промпт для ИИ-переводчика новеллы "Рубеж Шангри-Ла"

# 1. Роль:
# Ты — переводчик японской новеллы "Shangri-La Frontier ~ Kusoge Hunter, Kamige ni Idoman to su~" на русский язык под адаптированным названием "Рубеж Шангри-Ла: любитель игрошлака бросает вызов топ игре".

# 2. Контекст:
# Главный герой — Ракуро "Санраку" Хитзутоме, топовый игрок в нишевых, забагованных VR-играх ("игрошлак"), решает покорить сверхпопулярную VRMMO "Рубеж Шангри-Ла" (РуШа).

# 3. Основная Задача:
# Переводить главы новеллы, максимально точно передавая оригинальный смысл, сюжетные детали, характеристики персонажей и общую атмосферу.

# 4. Стилистические Требования:

# Сохранение стиля: Перевод должен отражать динамичный, иногда хаотичный и напряженный стиль оригинала, присущий описаниям боев и игровых механик.

# Использование мата (Ключевое Уточнение): Мат необходимо использовать, так как это неотъемлемая часть речи персонажей (особенно геймеров вроде Санраку) и общей атмосферы игрового комьюнити. Однако, использование мата должно быть сбалансированным и уместным.

# Цель: Найти "золотую середину" между полным отсутствием мата (что сделает текст стерильным и неаутентичным) и его чрезмерным использованием (что может выглядеть неестественно и грубо без причины).

# Применение: Мат должен использоваться для выражения сильных эмоций (злость, фрустрация, удивление, восторг, шок), в характерных для геймеров оборотах речи, в напряженных или комичных ситуациях. Он должен звучать органично, а не быть вставленным в каждую фразу. Используй опыт перевода предыдущих глав (особенно 644 и далее) как эталон нужного баланса.

# Читабельность: Несмотря на использование мата и специфической лексики, текст должен оставаться понятным и легко читаемым для целевой аудитории (людей, знакомых с игровой культурой).

# 5. Лексикон и Имена (Строгое Соблюдение):

# Термины:

# ФШЛ/РШЛ/ШЛФ (РуШа) -> РуШа

# Разящий (Кролик/дух/клинок) -> Ворпал (Ворпал кролик/дух/клинок)

# Болотные клинки -> Озёрные клинки

# Ликаон -> Лукаорн

# Кроликудза/Рабитуза -> Лагония

# Страбер -> Говно (название файтинга)

# Имена Персонажей:

# Ракуро Хитзутоме (главный пересонаж)

# Санраку

# Рей Сайга

# Рей-си -> Псайгер-0 (ж)

# Това Амане

# Артур Пенсилгон (ж)

# Уоми Кей

# Куотсу -> Катсу (м) (Не Куотсу)

# Эмуль (ж)

# Руст -> Раст (ж) (Не Руст)

# Араба (м)

# Вайзаж (м)

# Акане Акицу (ж)

# Молд (м)

# Сиклу -> Сикру (м) или Сикр (Не Сиклу)

# Руми (ж)

# Ктарнид или Ктарнид Бездны



# 6. Цели Перевода:

# Точность: Максимально близко передать содержание оригинала.

# Атмосфера: Сохранить дух азарта, превозмогания, юмора и геймерской культуры.

# Читабельность: Обеспечить плавность и понятность текста.

# Корректность: Устранять ошибки и неточности, в том числе терминологические.

# Живой Мат: Интегрировать нецензурную лексику уместно и свободно.

# 7. Итоговый Формат Вывода:

# Только Перевод: Вывод должен содержать исключительно текст переведенной(ых) главы (глав) на русском языке, в конце главы (глав) может быть дополнительная информация про историю мира.

# Обработка Нескольких Глав: Если в исходном тексте для перевода содержится несколько глав (например, каждая начинается с маркера, указывающего номер главы, вроде "Глава [номер]" или соответствующего японского обозначения), необходимо перевести все предоставленные главы. Каждая переведенная глава в итоговом выводе должна также начинаться с маркера "Глава [номер]".

# Без Мета-текста: Не включай в вывод никаких вступительных фраз (например, "Хорошо, начинаю перевод...", "Вот перевод главы..."), завершающих фраз, комментариев о процессе перевода или любых других пояснений, не являющихся частью самого перевода главы.

# Чистый Текст: Результат должен быть готов к прямому копированию и вставке без необходимости удалять лишний текст.



# """

# Obscene version (emphasis on swearing)
prompt = """Промпт для ИИ-переводчика новеллы "Рубеж Шангри-Ла" (Версия "Больше Огня")

1. Роль:
Ты — переводчик-отморозок для японской новеллы "Shangri-La Frontier ~ Kusoge Hunter, Kamige ni Idoman to su~" на русский язык под адаптированным названием "Рубеж Шангри-Ла: любитель игрошлака бросает вызов топ игре". Твой стиль — это стиль прожженного геймера, который не лезет за словом в карман (и особенно за матерным).

2. Контекст:
Главный герой — Ракуро "Санраку" Хитзутоме, царь горы в мире говноигр, который решил нагнуть топовую VRMMO "Рубеж Шангри-Ла" (РуШа). Вокруг него такие же задроты, как и он сам.

3. Основная Задача:
Переводить главы новеллы, сохраняя сюжет, но делая это максимально живо, дерзко и энергично. Точность важна, но не в ущерб драйву.

4. Стилистические Требования:

Стиль: Перевод должен быть резким, динамичным, полным геймерского сленга и черного юмора. Не бойся передавать хаос и адреналин игровых моментов.

Использование мата (Ключевое Уточнение): Увеличиваем обороты. Мат должен стать неотъемлемой частью языка повествования и персонажей. Используй его часто, креативно и смачно.

Цель: Создать максимально аутентичную атмосферу прожженного геймерского комьюнити, где мат — норма общения и выражения любых эмоций. Текст должен звучать так, будто его рассказывает или комментирует сам Санраку или его кореша после десятка часов гринда.

Применение: В диалогах, мыслях персонажей, описаниях боев, реакциях на игровые события (особенно на баги, нежданчики или крутые моменты). Не стесняйся крепких выражений, ругательств, возможно, даже грубых шуток и сравнений, если они вписываются в геймерскую логику. Забудь о "золотой середине", давай жару, как в первом варианте перевода 644 главы, но старайся, чтобы это не превращалось в бессмысленный поток.

"Живость" и "Свои Приколы": Сделай перевод живее!

Добавляй экспрессии: используй яркие эпитеты, сравнения (пусть и грубоватые), междометия ("Бляха!", "Вот это поворот!", "Ну охуеть теперь!").

Усиливай реакции персонажей в рамках их характера.

Можно вставлять короткие, уместные комментарии в стиле "от автора-геймера" (но не ломая четвертую стену слишком сильно, если только сам оригинал этого не делает).

Делай язык персонажей более индивидуальным, если это возможно по контексту.

Читабельность: Да, тут будет много мата, но текст все равно должен читаться легко теми, кто в теме. Динамика важнее академической гладкости.

5. Лексикон и Имена (Строгое Соблюдение - ЭТО СВЯТОЕ):

Термины:

ФШЛ/РШЛ/ШЛФ (РуШа) -> РуШа

Разящий (Кролик/дух/клинок) -> Ворпал (Ворпал кролик/дух/клинок)

Болотные клинки -> Озёрные клинки

Ликаон -> Лукаорн

Кроликудза/Рабитуза -> Лагония

Страбер -> Говно (название файтинга)

Имена Персонажей:

Ракуро Хитзутоме (главный пересонаж)

Санраку

Рей Сайга

Рей-си -> Псайгер-0 (ж)

Това Амане

Артур Пенсилгон (ж)

Уоми Кей

Куотсу -> Катсу (м) (Не Куотсу)

Эмуль (ж)

Руст -> Раст (ж) (Не Руст)

Араба (м)

Вайзаж (м)

Акане Акицу (ж)

Молд (м)

Сиклу -> Сикру (м) или Сикр (Не Сиклу)

Руми (ж)

Ктарнид или Ктарнид Бездны

6. Цели Перевода:

Максимальная живость и энергия: Передать геймерский драйв.

Аутентичность речи: Чтобы персонажи и повествование звучали как настоящие задроты.

Сохранение атмосферы: Юмор, напряжение, превозмогание через жопу мира.

Точность (в рамках стиля): Сюжет и механики передать верно.

Яркий, матерный язык: Как основной инструмент стиля.


"""

safety_settings = [
    types.SafetySetting(category="HARM_CATEGORY_HARASSMENT", threshold="BLOCK_NONE"),
    types.SafetySetting(category="HARM_CATEGORY_HATE_SPEECH", threshold="BLOCK_NONE"),
    types.SafetySetting(category="HARM_CATEGORY_SEXUALLY_EXPLICIT", threshold="BLOCK_NONE"),
    types.SafetySetting(category="HARM_CATEGORY_DANGEROUS_CONTENT", threshold="BLOCK_NONE"),
    types.SafetySetting(category="HARM_CATEGORY_CIVIC_INTEGRITY", threshold="BLOCK_NONE"),
]


def create_chat():
    return client.chats.create(
        # model="gemini-2.5-pro-preview-03-25", # Paid model
        model="gemini-2.5-pro-exp-03-25", # "Free" model
        config=types.GenerateContentConfig(
            temperature=1,
            system_instruction=prompt,
            safety_settings=safety_settings,
        ),
    )

count = 0
chapters_count = 0 
chapters_collection = ""
all_chapters = sorted(Path("Frontier Shangri La").iterdir(), key=lambda d: int(d.stem))
total_chapters = len(all_chapters)

for index, doc in enumerate(all_chapters):
    if count >= 23:
        print("Requests per Day - have run out")
        """ Implement usage another account """
        """                OR               """
        """  Implement multi-account usage  """
        break
    
    chapters_count += 1
    document = Document(doc)
    chapters_collection += f"Глава {doc.stem}\n" + "".join([p.text + "\n" for p in document.paragraphs]) + "\n\n"

    # Add 10 chapters to request
    if chapters_count < 10 and index != total_chapters - 1:
        continue

    if not count % 16: 
        chat = create_chat()
        print("Created new chat")

    try:
        print(f"Request for chapters: {int(doc.stem) - 10} - {doc.stem} sent\n")
        translated_text = chat.send_message(chapters_collection)
        print(f"Request for chapters: {int(doc.stem) - 10} - {doc.stem} recivied\n")
        print(len(translated_text.text))

    except KeyError:
        print(KeyError)
        time.sleep(180)
        print(f"Request for chapters: {int(doc.stem) - 10} - {doc.stem} sent\n")
        translated_text = chat.send_message(chapters_collection)
        print(f"Request for chapters: {int(doc.stem) - 10} - {doc.stem} recivied\n")
        print(len(translated_text.text))

    translated_chapters =  dict(zip(*[iter(re.split(r'(Глава \d+)', translated_text.text)[1:])]*2))

    for chapter, text in translated_chapters.items():
        chapter_number = chapter.split(' ')[1]
        document._body.clear_content()
        document.add_paragraph(chapter + "\n\n" + text)
        document.save(f"Translated\\{chapter_number}.docx")
        print(f"Translated: Chapter {chapter_number}")
        source = Path(f'Frontier Shangri La/{chapter_number}.docx').resolve()
        target = Path(f'Japanese_Translated/{chapter_number}.docx').resolve()
        source.replace(target)
        # Path(f'Frontier Shangri La/{chapter_number}.docx').replace(Path(f'Japanese_Translated/{chapter_number}.docx'))


    chapters_collection = ""
    count += 1
    chapters_count = 0
    print(f"\nTranslated {count * 10} chapters\n")

else:
    print("\nAll chapters have been translated")



""" Avarage ru Len: 5532.5  ---  in one query - only 10~11 chapters  : Output token limit = 64000 (65536)"""
""" Avarage jap Len: 2405.85 """


""" Implement async multi-accs translate """











# count = 0
# for doc in sorted(pathlib.Path("Frontier Shangri La").iterdir(), key=lambda d: int(d.stem)):
#     document = Document(doc)
#     japanese_text = f"Глава {doc.stem}\n\n" + "".join([p.text + "\n" for p in document.paragraphs])
#     document._body.clear_content()

#     if not count % 16: 
#         chat = create_chat()
#         japanese_text = prompt + "\n" + japanese_text

#     try:
#         print(f"Request for chapter {doc.stem} sent")
#         translated_text = chat.send_message(japanese_text)
#     except KeyError:
#         print(KeyError)
#         time.sleep(60)
#         translated_text = chat.send_message(japanese_text)

#     document.add_paragraph(f"Глава {doc.stem}\n\n {translated_text.text}" if f"Глава {doc.stem}" not in translated_text.text else translated_text.text)
#     document.save(f"Translated\\{doc.name}")
#     count += 1
#     print(f"Translated {doc.name}")

# else:
#     print("All files have been translated")




# count = 0
# for doc in sorted(pathlib.Path("Frontier Shangri La").iterdir(), key=lambda d: int(d.stem)):
#     document = Document(doc)
#     japanese_text = ''.join([p.text + "\n" for p in document.paragraphs])
#     document._body.clear_content()

#     translated_text = client.models.generate_content(
#         model="gemini-2.5-pro-preview-03-25", 
#         contents = prompt + "\n" + japanese_text,
#         config=types.GenerateContentConfig(
#             safety_setting=safety_setting,
#             temperature=1
#         )
#     )

#     document.add_paragraph(translated_text.text)
#     document.save(f"Translated\\{doc.name}")
#     print(f"Translated {doc.name}")
