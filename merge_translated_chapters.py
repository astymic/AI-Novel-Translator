from docxcompose.composer import Composer
from docx import Document 
from pathlib import Path

def merge_chapters():
    docs = sorted(Path("Translated").iterdir(), key=lambda x: int(x.stem))

    master = Document(docs[0])
    docs = docs[1:]

    composer = Composer(master)

    for doc in docs:
        page_break = Document()

        page_break.add_page_break()

        composer.append(page_break)

        next_doc = Document(doc)

        composer.append(next_doc)

    composer.save("Frontier Shangri La.docx")

merge_chapters()